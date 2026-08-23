# -*- coding: utf-8 -*-
"""
Экспорт двуязычной HTML-страницы контура «Радоновый риск» в .docx для правки
автором — без выполнения JS: разбирает исходник статически (BeautifulSoup)
и воспроизводит то же самое, что делает applyLang()/CSS в браузере.

Рассчитан на страницы вида radon-rational-method-calculator/index.html:
  - var I18N = {ru:{...}, en:{...}};            — базовый словарь
  - var SIMPLE = {ru:{...}, en:{...}};           — необязательный оверлей
    (в JS накладывается поверх I18N, когда MODE==="simple")
  - элементы с data-i18n="КЛЮЧ" — их innerHTML заменяется словарём
  - <optgroup id="rlGroupXX"> — атрибут label обновляется отдельно (в JS —
    через og.label=..., т.к. data-i18n на optgroup стёр бы дочерние <option>)
  - CSS-правила "body.simple СЕЛЕКТОР{display:none|block}" и безусловные
    правила ".класс{display:none}" — видимость по режиму (sci/simple)
  - <header>, <section aria-label="...">, <footer> — верхнеуровневая структура
  - section[aria-label="calculator"] — поля формы сворачиваются в таблицу
    "поле — значение по умолчанию — пояснение" (результат вычисления НЕ
    воспроизводится — он не статический контент, а вывод calc())
  - section[aria-label="flowchart"] — inline SVG не рендерится как картинка;
    экспортируется как текстовый список подписей узлов схемы (реальная схема
    остаётся только на живой странице — см. ссылку в шапке экспорта)
  - section[aria-label="table1"] — настоящая <table> конвертируется в таблицу
    Word как есть

Использование:
    python html_to_docx.py INDEX.HTML --mode sci   --lang ru -o out_sci.docx \
        --source-url "https://.../"
    python html_to_docx.py INDEX.HTML --mode simple --lang ru -o out_simple.docx \
        --source-url "https://.../simple.html"

Если страница устроена иначе (нет I18N/SIMPLE, другая структура секций) —
скрипт не подойдёт без правки; он не претендует на универсальный HTML→docx.
"""
import argparse
import re
import sys
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ───────────────────────── JS-словари: извлечение ─────────────────────────

def _find_matching_brace(text, open_idx):
    """open_idx указывает на '{'. Возвращает индекс парной '}', пропуская
    скобки внутри "..."-строк (с учётом \\-экранирования)."""
    depth = 0
    i = open_idx
    in_str = False
    while i < len(text):
        c = text[i]
        if in_str:
            if c == "\\":
                i += 2
                continue
            if c == '"':
                in_str = False
        else:
            if c == '"':
                in_str = True
            elif c == "{":
                depth += 1
            elif c == "}":
                depth -= 1
                if depth == 0:
                    return i
        i += 1
    raise ValueError("не нашлась парная скобка")


def extract_js_dict_block(source, var_name):
    """Возвращает сырой текст `{...}` для `var VAR_NAME = {...};`, либо None."""
    m = re.search(r"var\s+" + re.escape(var_name) + r"\s*=\s*\{", source)
    if not m:
        return None
    open_idx = m.end() - 1
    close_idx = _find_matching_brace(source, open_idx)
    return source[open_idx:close_idx + 1]


def extract_lang_subdicts(block_text):
    """Из '{ru:{...}, en:{...}}' возвращает {'ru': {...}, 'en': {...}}."""
    out = {}
    for lang in ("ru", "en"):
        m = re.search(r"\b" + lang + r"\s*:\s*\{", block_text)
        if not m:
            continue
        open_idx = m.end() - 1
        close_idx = _find_matching_brace(block_text, open_idx)
        inner = block_text[open_idx + 1:close_idx]
        pairs = re.findall(r'(\w+)\s*:\s*"((?:[^"\\]|\\.)*)"', inner)
        d = {}
        for key, val in pairs:
            d[key] = val.replace('\\"', '"').replace("\\n", " ")
        out[lang] = d
    return out


def load_dicts(source):
    i18n_block = extract_js_dict_block(source, "I18N")
    if i18n_block is None:
        raise SystemExit("var I18N = {...} не найден — этот скрипт не подходит для данной страницы")
    i18n = extract_lang_subdicts(i18n_block)
    simple_block = extract_js_dict_block(source, "SIMPLE")
    simple = extract_lang_subdicts(simple_block) if simple_block else {}
    return i18n, simple


# ───────────────────────── CSS: правила видимости ─────────────────────────

def extract_visibility_rules(source):
    """Возвращает (default_hide, default_show, simple_hide, simple_show) —
    множества CSS-селекторов из <style>, по одному правилу на строку вида
    'СЕЛЕКТОР(,СЕЛЕКТОР)*{display:none|block}', с необязательным префиксом
    'body.simple ' у каждой ветки через запятую."""
    style_m = re.search(r"<style[^>]*>(.*?)</style>", source, re.S)
    style_text = style_m.group(1) if style_m else ""
    default_hide, default_show = set(), set()
    simple_hide, simple_show = set(), set()
    rule_re = re.compile(r"([^{}]+)\{\s*display\s*:\s*(none|block)\s*;?\s*\}")
    for m in rule_re.finditer(style_text):
        disp = m.group(2)
        for sel in m.group(1).split(","):
            sel = sel.strip()
            if not sel or "display" in sel.lower():
                continue
            if sel.startswith("body.simple"):
                inner = sel[len("body.simple"):].strip()
                (simple_hide if disp == "none" else simple_show).add(inner)
            else:
                (default_hide if disp == "none" else default_show).add(sel)
    return default_hide, default_show, simple_hide, simple_show


def hidden_selectors_for_mode(rules, mode):
    default_hide, default_show, simple_hide, simple_show = rules
    hidden = set(default_hide)
    if mode == "simple":
        hidden -= simple_show
        hidden |= simple_hide
    else:
        hidden -= default_show
    return hidden


# ───────────────────────── применение словаря к DOM ─────────────────────────

def apply_dict(soup, lang_dict):
    for el in soup.select("[data-i18n]"):
        key = el.get("data-i18n")
        if key in lang_dict:
            el.clear()
            frag = BeautifulSoup(lang_dict[key], "html.parser")
            for child in list(frag.contents):
                el.append(child.extract())


def apply_optgroup_labels(soup, lang_dict):
    for code in ("RF", "WHO", "EPA", "EU"):
        key = "rlGroup" + code
        og = soup.find(id="rlGroup" + code)
        if og is not None and key in lang_dict:
            og["label"] = lang_dict[key]


def prune_hidden(soup, hidden_selectors):
    for sel in hidden_selectors:
        try:
            for el in soup.select(sel):
                if el.parent is not None:
                    el.decompose()
        except Exception:
            pass  # неподдерживаемый селектор — пропускаем, не валим экспорт


def strip_eroa_in_simple(soup, mode):
    if mode == "simple":
        opt = soup.select_one('#ctUnit option[value="eroa"]')
        if opt is not None:
            opt.decompose()


# ───────────────────────── docx: вставка hyperlink ─────────────────────────

def add_hyperlink(paragraph, url, text, color="1155CC", underline=True):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rpr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ───────────────────────── docx: инлайн-разметка ─────────────────────────

def add_inline(paragraph, node, bold=False, italic=False, sub=False, sup=False):
    if isinstance(node, NavigableString):
        text = str(node)
        if text == "":
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if sub:
            run.font.subscript = True
        if sup:
            run.font.superscript = True
        return
    if not isinstance(node, Tag):
        return
    name = node.name
    if name == "br":
        paragraph.add_run().add_break()
        return
    if name == "a" and node.get("href"):
        add_hyperlink(paragraph, node["href"], node.get_text())
        return
    nb = bold or name in ("b", "strong")
    ni = italic or name in ("i", "em", "cite")
    nsub = sub or name == "sub"
    nsup = sup or name == "sup"
    for child in node.children:
        add_inline(paragraph, child, nb, ni, nsub, nsup)


def add_para(doc, node, style=None, base_italic=False):
    p = doc.add_paragraph(style=style)
    for child in node.children:
        add_inline(p, child, italic=base_italic)
    return p


# ───────────────────────── docx: блок-цитата ─────────────────────────

def add_quote(doc, bq_tag):
    cite = bq_tag.find("cite")
    cite_text = cite.get_text(strip=True) if cite else ""
    if cite:
        cite.extract()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.space_after = Pt(2)
    for child in bq_tag.children:
        add_inline(p, child, italic=True)
    for run in p.runs:
        run.font.size = Pt(10.5)
    if cite_text:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.35)
        p2.paragraph_format.space_after = Pt(10)
        r = p2.add_run(cite_text)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


# ───────────────────────── docx: обход общего блока ─────────────────────────

BLOCK_TAGS = {"h1", "h2", "h3", "h4", "p", "blockquote", "ul", "ol", "table", "div", "details"}


def has_block_children(tag):
    return any(isinstance(c, Tag) and c.name in BLOCK_TAGS for c in tag.children)


def handle_node(doc, node):
    """Диспетчер по одному узлу: h1/h2/h3 → заголовки, p → абзац,
    blockquote.quote → цитата, ul/ol → списки, table → таблица,
    details → summary+тело, div/section/header/footer → рекурсивно в детей."""
    if not isinstance(node, Tag):
        return
    name = node.name
    classes = node.get("class") or []
    if name == "h1":
        add_para(doc, node, style="Title")
    elif name == "h2":
        add_para(doc, node, style="Heading 1")
    elif name == "h3":
        add_para(doc, node, style="Heading 2")
    elif name == "h4":
        add_para(doc, node, style="Heading 3")
    elif name == "blockquote" and "quote" in classes:
        add_quote(doc, node)
    elif name == "p":
        add_para(doc, node)
    elif name in ("ul", "ol"):
        style = "List Bullet" if name == "ul" else "List Number"
        for li in node.find_all("li", recursive=False):
            if has_block_children(li):
                # <li> с блочным содержимым (напр. ol.steps: h3+p+blockquote) —
                # раскрываем как обычные блоки, а не сплющиваем в один абзац
                walk_block(doc, li)
            else:
                add_para(doc, li, style=style)
    elif name == "table":
        add_table(doc, node)
    elif name == "details":
        summary = node.find("summary", recursive=False)
        if summary:
            add_para(doc, summary, style="Heading 3")
        for child in node.children:
            if isinstance(child, Tag) and child.name != "summary":
                handle_node(doc, child)
    elif name in ("div", "section", "header", "footer"):
        # div с чисто инлайновым содержимым (напр. резолвленная data-i18n-формула
        # с <br>/<sub> прямо в детях) — как один абзац, иначе блоки молча теряются
        if name == "div" and not has_block_children(node):
            add_para(doc, node)
        else:
            walk_block(doc, node)
    # svg / script / style / form controls — сознательно пропускаются здесь;
    # калькулятор и схема обрабатываются отдельными специализированными
    # функциями (см. add_calculator_section / add_flowchart_section)


def walk_block(doc, container):
    """Обходит container.children и диспетчеризует каждый через handle_node."""
    for node in container.children:
        if isinstance(node, NavigableString):
            continue
        handle_node(doc, node)


def add_table(doc, table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    ncols = max(len(r.find_all(["td", "th"])) for r in rows)
    t = doc.add_table(rows=0, cols=ncols)
    t.style = "Light Grid Accent 1"
    for r in rows:
        cells = r.find_all(["td", "th"])
        row = t.add_row()
        for i, c in enumerate(cells):
            if i >= ncols:
                break
            cell_p = row.cells[i].paragraphs[0]
            for child in c.children:
                add_inline(cell_p, child, bold=(c.name == "th"))


# ───────────────────────── docx: секция «калькулятор» ─────────────────────────

def describe_field(doc, field_div):
    label = field_div.find(["label"])
    hint = field_div.find("p", class_="hint")
    inputs = field_div.select("input, select")
    if not inputs and not label:
        return
    p = doc.add_paragraph()
    if label is not None:
        for child in label.children:
            add_inline(p, child, bold=True)
    else:
        p.add_run("(поле)").bold = True

    for inp in inputs:
        if inp.name == "input":
            val = inp.get("value", "")
            doc.add_paragraph("Значение по умолчанию: " + val, style="List Bullet")
        elif inp.name == "select":
            opts = inp.find_all("option")
            lines = []
            for og in inp.find_all("optgroup"):
                lines.append("[" + og.get("label", "") + "]")
                for o in og.find_all("option"):
                    mark = " ← по умолчанию" if o.has_attr("selected") else ""
                    lines.append("  " + o.get_text(" ", strip=True) + mark)
            if not inp.find_all("optgroup"):
                for o in opts:
                    mark = " ← по умолчанию" if o.has_attr("selected") else ""
                    lines.append(o.get_text(" ", strip=True) + mark)
            for ln in lines:
                doc.add_paragraph(ln, style="List Bullet")
    if hint:
        hp = doc.add_paragraph()
        hp.paragraph_format.space_after = Pt(10)
        for child in hint.children:
            add_inline(hp, child, italic=True)
        for run in hp.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)


def add_calculator_section(doc, section_tag):
    for node in section_tag.children:
        if not isinstance(node, Tag):
            continue
        if node.name == "h2":
            add_para(doc, node, style="Heading 1")
        elif node.name == "p" and "hint" in (node.get("class") or []):
            add_para(doc, node)
        elif node.name == "div" and "grid" in (node.get("class") or []):
            for field_div in node.find_all("div", recursive=False):
                describe_field(doc, field_div)
        elif node.name == "details":
            summary = node.find("summary", recursive=False)
            if summary:
                add_para(doc, summary, style="Heading 3")
            body = node.find("div", class_="refbody")
            if body:
                walk_block(doc, body)
        elif node.name == "div" and node.get("id") == "out":
            continue  # результат — вывод calc(), не статический контент
        elif node.name in ("h3", "h4", "p", "ul", "ol"):
            handle_node(doc, node)
    doc.add_paragraph(
        "Итоговый вывод и промежуточные величины (запасы, пороги, чипы) "
        "вычисляются на живой странице калькулятора и здесь не воспроизводятся."
    ).runs[0].italic = True


# ───────────────────────── docx: секция «схема» ─────────────────────────

def add_flowchart_section(doc, section_tag, live_url=None):
    for node in section_tag.find_all(["h2", "h3", "p"], recursive=False):
        style = "Heading 1" if node.name == "h2" else ("Heading 2" if node.name == "h3" else None)
        add_para(doc, node, style=style)
    legend = section_tag.find("ul", class_="hint-list")
    if legend:
        for li in legend.find_all("li", recursive=False):
            add_para(doc, li, style="List Bullet")
    svg = section_tag.find("svg")
    if svg:
        doc.add_paragraph(
            "Текст блок-схемы (сама диаграмма — векторная, на живой странице" +
            (f": {live_url}" if live_url else "") + "):"
        ).runs[0].italic = True
        for t in svg.find_all("text"):
            txt = t.get_text(" ", strip=True)
            if txt:
                doc.add_paragraph(txt, style="List Bullet")


# ───────────────────────── основной сценарий ─────────────────────────

def build_docx(html_path, mode, lang, out_path, source_url=None):
    source = open(html_path, encoding="utf-8").read()
    i18n, simple = load_dicts(source)
    lang_dict = dict(i18n.get(lang, {}))
    if mode == "simple":
        lang_dict.update(simple.get(lang, {}))

    soup = BeautifulSoup(source, "html.parser")
    apply_dict(soup, lang_dict)
    apply_optgroup_labels(soup, lang_dict)
    strip_eroa_in_simple(soup, mode)

    rules = extract_visibility_rules(source)
    hidden = hidden_selectors_for_mode(rules, mode)
    prune_hidden(soup, hidden)

    doc = Document()
    if source_url:
        note = doc.add_paragraph()
        note.add_run(f"Экспортировано из {source_url} — для правки текста автором. "
                      "Форматирование (заголовки, цитаты, списки) сохранено; "
                      "интерактивный калькулятор описан таблицей полей; "
                      "векторная схема — текстовым списком подписей.").italic = True
        doc.add_paragraph()

    header = soup.find("header")
    if header:
        walk_block(doc, header)

    for aria in ("introduction", "steps"):
        sec = soup.find("section", attrs={"aria-label": aria})
        if sec:
            walk_block(doc, sec)

    fc = soup.find("section", attrs={"aria-label": "flowchart"})
    if fc:
        add_flowchart_section(doc, fc, live_url=source_url)

    calc = soup.find("section", attrs={"aria-label": "calculator"})
    if calc:
        add_calculator_section(doc, calc)

    t1 = soup.find("section", attrs={"aria-label": "table1"})
    if t1:
        walk_block(doc, t1)

    footer = soup.find("footer")
    if footer:
        walk_block(doc, footer)

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("html_file")
    ap.add_argument("--mode", choices=["sci", "simple"], default="sci")
    ap.add_argument("--lang", choices=["ru", "en"], default="ru")
    ap.add_argument("-o", "--out", required=True)
    ap.add_argument("--source-url", default=None)
    args = ap.parse_args()
    path = build_docx(args.html_file, args.mode, args.lang, args.out, args.source_url)
    print("OK:", path)


if __name__ == "__main__":
    main()
